# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document


DOC_DIR = Path(__file__).resolve().parent / "Architecture"


def main():
    docx_files = sorted(p for p in DOC_DIR.glob("*.docx") if not p.name.startswith("~$"))
    if not docx_files:
        raise SystemExit("No DOCX files found")

    for path in docx_files:
        doc = Document(path)
        paragraph_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if len(paragraph_text) < 8:
            raise SystemExit(f"{path.name}: not enough paragraphs")
        if len(doc.tables) < 1:
            raise SystemExit(f"{path.name}: metadata table is missing")
        required = ["Карта сущности", "Ответственность", "Правила сопровождения", "Что проверять при изменении"]
        text = "\n".join(paragraph_text)
        for marker in required:
            if marker not in text:
                raise SystemExit(f"{path.name}: missing section {marker}")
        print(f"OK {path.name}: paragraphs={len(paragraph_text)} tables={len(doc.tables)}")


if __name__ == "__main__":
    main()
